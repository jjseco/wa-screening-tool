"""
Generate a professional Word document screening report.
"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import OUTPUTS_SUMMARIES
from scripts.registry import LAYER_REGISTRY
from scripts.project_types import PROJECT_TYPES


def _parse_dataset(source):
    """Extract dataset code from a source string such as 'Official WA - DWER-046'."""
    if " - " in source:
        return source.split(" - ", 1)[1]
    return source


def _get_provider(source):
    """Extract provider name from source string such as 'Official WA - DWER-046'."""
    if " - " in source:
        return source.split(" - ", 1)[0]
    return source


def _get_citation(layer_key, entry):
    """Build a formatted citation string for a registry layer."""
    label = entry.get("label", layer_key)
    code = entry.get("dataset_code", "")
    source = entry.get("source", "")
    dataset_year = entry.get("dataset_year", 2024)
    accessed_year = entry.get("accessed_year", 2026)

    if "OSM" in source:
        return (
            f"OpenStreetMap contributors ({dataset_year}). {label}. "
            f"Geofabrik Australia extract. Accessed March {accessed_year} from download.geofabrik.de"
        )
    provider = _get_provider(source) if source else ""
    if provider:
        return (
            f"{provider} ({dataset_year}). {label} [{code}]. "
            f"Accessed March {accessed_year} from data.wa.gov.au"
        )
    return ""


def _set_col_bold(table, col_idx):
    """Bold all paragraphs in a table column (used for header row)."""
    row = table.rows[0]
    cell = row.cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _add_table(doc, headers, rows):
    """
    Add a Table Grid table with a bold centred header row.
    Returns the table object.
    """
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    # Header row
    hdr_row = table.rows[0]
    for i, heading in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = heading
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tbl_row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            tbl_row.cells[c_idx].text = str(val) if val is not None else ""

    return table


def _summary_paragraph(theme_summary, site_name, latitude, longitude):
    """Build the executive summary paragraph text."""
    high = [t for t, ts in theme_summary.items() if ts["rating"] == "HIGH"]
    med = [t for t, ts in theme_summary.items() if ts["rating"] == "MEDIUM"]
    low = [t for t, ts in theme_summary.items() if ts["rating"] == "LOW"]

    intro = (
        f"A preliminary desktop environmental screening was conducted for {site_name} "
        f"at coordinates {latitude:.6f}\u00b0, {longitude:.6f}\u00b0. "
    )
    parts = []
    if high:
        parts.append(f"HIGH risk was identified across {len(high)} theme(s): {', '.join(high)}.")
    if med:
        parts.append(f"MEDIUM risk was identified across {len(med)} theme(s): {', '.join(med)}.")
    if low:
        parts.append(f"LOW risk was identified across {len(low)} theme(s).")
    if not parts:
        parts.append("No significant environmental sensitivities were identified within the screening area.")

    closing = (
        " Further investigation and specialist assessment is recommended for all themes "
        "rated MEDIUM or HIGH. Results should be verified against current authoritative "
        "sources before use in formal environmental assessments."
    )
    return intro + " ".join(parts) + closing


def generate_word_report(
    results,
    site_name,
    latitude,
    longitude,
    project_type,
    buffer_distances,
    theme_summary,
    layers,
    buffers=None,
    selected_themes=None,
    session_id="",
):
    """
    Generate a professional Word document screening report.

    Returns the file path of the saved .docx file.
    """
    doc = Document()

    # Default font
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # ----------------------------------------------------------------
    # COVER SECTION
    # ----------------------------------------------------------------
    title_para = doc.add_heading("Preliminary Environmental Screening Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    def _info_line(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(str(value))

    _info_line("Site Name", site_name)
    _info_line("Coordinates", f"{latitude:.6f}\u00b0 N, {longitude:.6f}\u00b0 E")
    _info_line("Date", date.today().strftime("%d %B %Y"))
    _info_line("Project Type", project_type)
    _info_line("Buffer Distances", " / ".join(f"{d}m" for d in buffer_distances))

    doc.add_paragraph()

    # ----------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # ----------------------------------------------------------------
    doc.add_heading("Executive Summary", level=1)

    themes_ordered = selected_themes if selected_themes else list(theme_summary.keys())
    risk_rows = []
    for theme in themes_ordered:
        if theme in theme_summary:
            ts = theme_summary[theme]
            risk_rows.append([theme, ts["rating"], ts["score"], ts["key_finding"]])

    if risk_rows:
        _add_table(doc, ["Theme", "Risk Rating", "Score", "Key Finding"], risk_rows)
        doc.add_paragraph()

    doc.add_paragraph(_summary_paragraph(theme_summary, site_name, latitude, longitude))
    doc.add_paragraph()

    # ----------------------------------------------------------------
    # SCREENING RESULTS BY THEME
    # ----------------------------------------------------------------
    doc.add_heading("Screening Results by Theme", level=1)

    layer_weights = {}
    if project_type in PROJECT_TYPES:
        layer_weights = PROJECT_TYPES[project_type].get("layer_weights", {})

    themes_in_results = list(dict.fromkeys(r["theme"] for r in results))
    for theme in themes_in_results:
        theme_results = [r for r in results if r["theme"] == theme]
        # Skip layers with weight == 0
        visible = [r for r in theme_results if layer_weights.get(r["layer_key"], 1) != 0]
        if not visible:
            continue

        doc.add_heading(theme, level=2)

        tbl_headers = [
            "Layer", "Source", "Present", "Max Relevance",
            "Primary Trigger", "Nearest (m)",
            "Detected Features", "Interpretation",
        ]
        tbl_rows = []
        for r in visible:
            nearest = r.get("nearest_distance_m")
            nearest_str = "Intersects site" if r["site_intersect"] else (str(nearest) if nearest else "\u2014")
            detected = ", ".join(r["nearby_names"]) if r["nearby_names"] else "\u2014"
            _e = LAYER_REGISTRY.get(r["layer_key"], {})
            source_str = f"{_e.get('label', r['layer'])} ({_e.get('dataset_code', '')})"
            tbl_rows.append([
                r["layer"],
                source_str,
                "Yes" if r["present"] else "No",
                r.get("max_relevance", "\u2014"),
                r.get("primary_trigger", "\u2014"),
                nearest_str,
                detected,
                r.get("interpretation", "\u2014"),
            ])

        _add_table(doc, tbl_headers, tbl_rows)
        doc.add_paragraph()

    # ----------------------------------------------------------------
    # SENSITIVE RECEPTORS DETAIL
    # ----------------------------------------------------------------
    receptor_keys = ["hospitals", "schools", "community_facilities"]
    has_receptors = any(k in layers for k in receptor_keys) or "cadastre" in layers

    if has_receptors and buffers:
        doc.add_heading("Sensitive Receptors Detail", level=1)

        # Derive site point as centroid of smallest buffer
        site_point = None
        try:
            smallest_dist = min(buffers.keys())
            site_point = buffers[smallest_dist].geometry.iloc[0].centroid
        except Exception:
            pass

        # Cadastre parcel counts
        if "cadastre" in layers:
            doc.add_heading("Cadastral Parcels", level=2)
            cadastre_gdf = layers["cadastre"]
            for dist in buffer_distances:
                if dist in buffers:
                    try:
                        buf_union = buffers[dist].geometry.union_all()
                        count = len(cadastre_gdf[cadastre_gdf.geometry.intersects(buf_union)])
                        doc.add_paragraph(f"Parcels within {dist}m: {count}")
                    except Exception:
                        pass

        # Individual receptor lists
        receptor_label_map = {
            "hospitals": "Hospitals",
            "schools": "Schools",
            "community_facilities": "Community Facilities",
        }

        max_dist = max(buffer_distances) if buffer_distances else max(buffers.keys())
        if max_dist not in buffers:
            max_dist = max(buffers.keys())

        for layer_key, label in receptor_label_map.items():
            if layer_key not in layers:
                continue
            try:
                buf_union = buffers[max_dist].geometry.union_all()
                gdf = layers[layer_key]
                nearby = gdf[gdf.geometry.intersects(buf_union)].copy()
                if nearby.empty:
                    continue

                doc.add_heading(label, level=2)
                receptor_rows = []
                for _, row in nearby.iterrows():
                    name = ""
                    if "name" in row and str(row["name"]) not in ["nan", "None", ""]:
                        name = str(row["name"])
                    if site_point is not None:
                        raw_dist = float(row.geometry.distance(site_point))
                        dist_str = "Intersects site" if raw_dist < 1 else f"{raw_dist:.0f}m"
                    else:
                        dist_str = "\u2014"
                    receptor_rows.append([name or "\u2014", dist_str])

                if receptor_rows:
                    _add_table(doc, ["Name", "Distance to Site"], receptor_rows)
                doc.add_paragraph()

            except Exception as e:
                print(f"Receptor detail error for {layer_key}: {e}")

    # ----------------------------------------------------------------
    # DATA SOURCES
    # ----------------------------------------------------------------
    doc.add_heading("Data Sources", level=1)

    seen_keys = set()
    source_rows = []
    for layer_key in (layers or {}):
        if layer_key in seen_keys:
            continue
        seen_keys.add(layer_key)
        entry = LAYER_REGISTRY.get(layer_key, {})
        label = entry.get("label", layer_key)
        code = entry.get("dataset_code", "")
        source = entry.get("source", "")
        theme = entry.get("theme", "")
        provider = _get_provider(source) if source else ""
        citation = _get_citation(layer_key, entry)
        source_rows.append([theme, label, code, provider, citation])

    if source_rows:
        source_rows.sort(key=lambda x: (x[0], x[1]))
        _add_table(doc, ["Theme", "Layer", "Dataset Code", "Provider", "Citation"], source_rows)

    doc.add_paragraph()

    # ----------------------------------------------------------------
    # DISCLAIMER
    # ----------------------------------------------------------------
    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        "This report is intended for preliminary desktop screening purposes only. "
        "Results are based on publicly available spatial datasets and should be verified "
        "against current authoritative sources before use in formal environmental assessments. "
        "This screening does not constitute a formal environmental assessment, environmental "
        "impact statement, or specialist technical report. Data currency, completeness, and "
        "accuracy cannot be guaranteed. Always consult a qualified environmental professional "
        "for project-specific advice. The report operator accepts no liability for decisions "
        "made solely on the basis of this screening report."
    )

    # ----------------------------------------------------------------
    # SAVE
    # ----------------------------------------------------------------
    out_dir = "/tmp/wa_outputs/reports"
    os.makedirs(out_dir, exist_ok=True)
    prefix = f"{session_id}_" if session_id else ""
    safe_name = site_name.replace(" ", "_").replace("/", "-")
    file_name = f"{prefix}{safe_name}_screening_report.docx"
    file_path = os.path.join(out_dir, file_name)
    doc.save(file_path)
    print(f"Word report saved to: {file_path}")
    return file_path
